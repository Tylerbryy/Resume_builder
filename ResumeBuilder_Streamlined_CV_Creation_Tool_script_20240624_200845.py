import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Constants
NAME = "Tyler Gibbs"
TITLE = "SOFTWARE ENGINEER"
CONTACT_INFO = {
    'phone': "405-320-8212",
    'email': "tylergibbs048@gmail.com",
    'location': "Blanchard, OK 73010\nAmerican Indian (Choctaw)",
    'linkedin': "linkedin.com/in/tylergibbss",
    'website': "tylergibbs.dev"
}
PROFESSIONAL_SUMMARY = (
    "Highly motivated university student with a concentration in software engineering, "
    "currently holding a 3.8 GPA. With more than 4 years of relevant work experience and a "
    "history of extracurricular involvement, I'm keen on channeling my problem-solving skills "
    "and passion for technology. Renowned for my results-oriented approach and team-centric "
    "mentality, I am adept at developing effective software solutions and fostering team success. "
    "I like to approach things from first principles, ensuring a deep understanding of the fundamentals "
    "to create innovative and efficient solutions."
)
EXPERIENCES = [
    {
        "title": "DATA ANALYST/SOFTWARE ENGINEER",
        "company": "Lexis Nexis",
        "date": "Aug 2023 - Present",
        "responsibilities": [
            "Performed data analysis using SQL, Python, and Excel.",
            "Created data reports and dashboards with Tableau and Power BI.",
            "Developed productivity-enhancing software tools.",
            "Saved $142,000 through efficient software solutions.",
            "Collaborated to identify key performance indicators.",
            "Conducted statistical analysis and predictive modeling.",
            "Managed data cleaning and validation."
        ]
    },
    {
        "title": "SOFTWARE ENGINEER FREELANCER",
        "company": "Relevant Technical Experience",
        "date": "Jul 2022 – Aug 2023",
        "responsibilities": [
            "Translated client needs into technical requirements.",
            "Developed and debugged software applications.",
            "Used Java, Python, and C++ for various projects.",
            "Adopted new technologies and frameworks.",
            "Applied Agile methodologies for project management."
        ]
    }
]
EDUCATION = {
    'degree': "BACHELOR OF SCIENCE IN COMPUTER SCIENCE CANDIDATE",
    'institution': "University Of Oklahoma, Norman, OK, US",
    'graduation': "Expected graduation: May 2025",
    'coursework': "Data Structures and Algorithms, Object-Oriented Design, Operating Systems, Data Analytics",
    'activities': "Member, University of Oklahoma Robotics Team"
}
PROJECTS = [
    {
        "title": "Backwork",
        "description": (
            "Automates the medical coding process using AI-powered medical billing code extraction from documents. "
            "Features include user authentication, profile management, admin dashboard, responsive web design using "
            "Tailwind CSS, and integration with Finetuned LLM for document processing."
        )
    },
    {
        "title": "RiskLabs",
        "description": (
            "Uses AI agents to analyze stock market trends and provide real-time insights. Features include real-time reports, "
            "comprehensive insights, strategic investment recommendations, and tailored analysis."
        )
    },
    {
        "title": "Zinbo",
        "description": (
            "Uses LLM function calling to clean your email inbox. Features include Gmail integration, terminal interface, Google Cloud "
            "Gmail API, and support for multiple AI models."
        )
    },
    {
        "title": "Knightsbridge Website",
        "description": "My software engineering firm website. Features include Next.js, Supabase, TailwindCSS, and various UI libraries."
    },
    {
        "title": "AI-Powered Web Scraper",
        "description": (
            "An AI-powered web scraper that allows you to describe what you want on the page, and it scrapes it for you in a well-formatted JSON."
        )
    }
]
SKILLS = {
    "Programming Languages": ["Python", "C++", "Java"],
    "Web Technologies": ["Django", "Flask", "HTML/CSS", "JavaScript"],
    "Data Analysis": ["SQL", "Pandas", "NumPy", "Data Visualization"],
    "Tools & Methodologies": ["Git", "Agile", "CI/CD", "Unit Testing"],
    "Concepts": ["Object-Oriented Design", "Data Structures", "Algorithms", "Problem Solving"]
}
VOLUNTEER_EXPERIENCES = [
    {
        "title": "CODING BOOTCAMP VOLUNTEER",
        "organization": "University Of Oklahoma",
        "location": "Norman, OKLAHOMA",
        "description": "Assisted in teaching coding basics to local school students, developing their interest in software engineering and space technologies."
    },
    {
        "title": "TUTOR",
        "organization": "University Of Oklahoma",
        "location": "Norman, OKLAHOMA",
        "description": "Volunteered to tutor students in Computer Science, emphasizing programming, algorithms, and data structures."
    }
]
CERTIFICATIONS = [
    {
        "name": "Machine Learning",
        "issuer": "Stanford University School of Engineering",
        "date": "Jun 2024",
        "credential_id": "NVQB57DCUVZ9"
    },
    {
        "name": "Robotics: Aerial Robotics",
        "issuer": "University of Pennsylvania",
        "date": "Jun 2024",
        "credential_id": "TZAB8NFLTH5V"
    },
    {
        "name": "Self-Driving Cars",
        "issuer": "University of Toronto",
        "date": "Jun 2024",
        "credential_id": "26XLUZUSNHEW"
    }
]

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4F81BD')
    pBdr.append(bottom)

def add_section_header(doc, text):
    header = doc.add_paragraph()
    run = header.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    add_horizontal_line(header)

def create_resume():
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Name and Title
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(NAME)
    name_run.bold = True
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(TITLE)
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    
    # Contact Info
    contact_table = doc.add_table(rows=1, cols=3)
    contact_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_table.cell(0, 0).text = f"{CONTACT_INFO['phone']}\n{CONTACT_INFO['email']}"
    contact_table.cell(0, 1).text = CONTACT_INFO['location']
    contact_table.cell(0, 2).text = f"{CONTACT_INFO['linkedin']}\n{CONTACT_INFO['website']}"
    
    for row in contact_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    # Add horizontal line
    add_horizontal_line(doc.add_paragraph())
    
    # Summary
    add_section_header(doc, "PROFESSIONAL SUMMARY")
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary.add_run(PROFESSIONAL_SUMMARY)
    summary_run.font.name = 'Calibri'
    summary_run.font.size = Pt(11)
    
    # Experience
    add_section_header(doc, "PROFESSIONAL EXPERIENCE")
    
    for exp in EXPERIENCES:
        p = doc.add_paragraph()
        title_run = p.add_run(exp['title'])
        title_run.bold = True
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        company_run = p.add_run(f"\n{exp['company']} – {exp['date']}")
        company_run.italic = True
        company_run.font.name = 'Calibri'
        company_run.font.size = Pt(11)
        for resp in exp['responsibilities']:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.paragraph_format.left_indent = Inches(0.25)
            resp_run = bullet.add_run(resp)
            resp_run.font.name = 'Calibri'
            resp_run.font.size = Pt(11)
    
    # Education
    add_section_header(doc, "EDUCATION")
    
    edu = doc.add_paragraph()
    edu_run = edu.add_run(EDUCATION['degree'])
    edu_run.bold = True
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(12)
    edu_run.font.color.rgb = RGBColor(0, 0, 0)
    edu_run = edu.add_run(f"\n{EDUCATION['institution']}")
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(11)
    edu_run = edu.add_run(f"\t\t\t\t{EDUCATION['graduation']}")
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(11)
    edu_run.italic = True
    
    coursework_header = doc.add_paragraph()
    coursework_run = coursework_header.add_run("Relevant Coursework:")
    coursework_run.bold = True
    coursework_run.font.name = 'Calibri'
    coursework_run.font.size = Pt(11)
    
    coursework = doc.add_paragraph(style='List Bullet')
    coursework.paragraph_format.left_indent = Inches(0.25)
    coursework_run = coursework.add_run(EDUCATION['coursework'])
    coursework_run.font.name = 'Calibri'
    coursework_run.font.size = Pt(11)
    
    activities_header = doc.add_paragraph()
    activities_run = activities_header.add_run("Extracurricular Activities:")
    activities_run.bold = True
    activities_run.font.name = 'Calibri'
    activities_run.font.size = Pt(11)
    
    activities = doc.add_paragraph(style='List Bullet')
    activities.paragraph_format.left_indent = Inches(0.25)
    activities_run = activities.add_run(EDUCATION['activities'])
    activities_run.font.name = 'Calibri'
    activities_run.font.size = Pt(11)
    
    # Projects
    add_section_header(doc, "PROJECTS")
    
    for project in PROJECTS:
        p = doc.add_paragraph()
        title_run = p.add_run(project['title'])
        title_run.bold = True
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        desc_run = p.add_run(f"\n{project['description']}")
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(11)
    
    # Skills
    add_section_header(doc, "TECHNICAL SKILLS")
    
    for category, skill_list in SKILLS.items():
        p = doc.add_paragraph()
        category_run = p.add_run(f"{category}: ")
        category_run.bold = True
        category_run.font.name = 'Calibri'
        category_run.font.size = Pt(11)
        
        skills_run = p.add_run(", ".join(skill_list))
        skills_run.font.name = 'Calibri'
        skills_run.font.size = Pt(11)
    
    # Volunteer Experience
    add_section_header(doc, "VOLUNTEER EXPERIENCE")
    
    for exp in VOLUNTEER_EXPERIENCES:
        p = doc.add_paragraph()
        title_run = p.add_run(exp['title'])
        title_run.bold = True
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        org_run = p.add_run(f"\n{exp['organization']} – {exp['location']}")
        org_run.italic = True
        org_run.font.name = 'Calibri'
        org_run.font.size = Pt(11)
        desc_run = p.add_run(f"\n{exp['description']}")
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(11)
    
    # Certifications
    add_section_header(doc, "CERTIFICATIONS")
    
    for cert in CERTIFICATIONS:
        p = doc.add_paragraph()
        name_run = p.add_run(cert['name'])
        name_run.bold = True
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(11)
        
        details_run = p.add_run(f"\n{cert['issuer']} • Issued {cert['date']} • Credential ID {cert['credential_id']}")
        details_run.font.name = 'Calibri'
        details_run.font.size = Pt(10)
        details_run.italic = True
        
        p.add_run("\n")
    
    doc.save('Tyler_Gibbs_Resume.docx')

if __name__ == "__main__":
    create_resume()
    print("Resume created successfully as 'Tyler_Gibbs_Resume.docx'")
